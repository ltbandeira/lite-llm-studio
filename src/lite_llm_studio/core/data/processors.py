"""
Module core.data.processors
---------------------------

This module contains data processing functionality for the LiteLLM Studio project.
"""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from dataclasses import dataclass
import tempfile
import os


@dataclass
class ProcessedDocument:
    """Represents a processed document with extracted content."""

    filename: str
    content: str
    metadata: Dict[str, Any]
    word_count: int
    file_type: str
    size_bytes: int
    processing_status: str = "success"
    error_message: Optional[str] = None


class DocumentProcessor:
    """
    Document processor using docling for various file formats.

    Supports processing of PDF, Word, PowerPoint, and text files.
    """

    def __init__(self, logger_name: str = "app.data.processor"):
        """Initialize the document processor."""
        self.logger = logging.getLogger(logger_name)
        self.supported_extensions = {".pdf", ".docx", ".doc", ".pptx", ".ppt", ".txt", ".md", ".rtf"}

    def is_supported_file(self, file_path: Union[str, Path]) -> bool:
        """Check if the file type is supported for processing."""
        path = Path(file_path)
        return path.suffix.lower() in self.supported_extensions

    def process_file(self, file_path: Union[str, Path]) -> ProcessedDocument:
        """
        Process a single file and extract its content.

        Args:
            file_path: Path to the file to process

        Returns:
            ProcessedDocument: Object containing processed content and metadata
        """
        path = Path(file_path)

        if not path.exists():
            return ProcessedDocument(
                filename=path.name,
                content="",
                metadata={},
                word_count=0,
                file_type="unknown",
                size_bytes=0,
                processing_status="error",
                error_message=f"File not found: {path}",
            )

        if not self.is_supported_file(path):
            return ProcessedDocument(
                filename=path.name,
                content="",
                metadata={},
                word_count=0,
                file_type=path.suffix.lower(),
                size_bytes=path.stat().st_size,
                processing_status="error",
                error_message=f"Unsupported file type: {path.suffix}",
            )

        try:
            self.logger.info(f"Processing file: {path.name}")

            # Get file stats
            file_stats = path.stat()
            file_size = file_stats.st_size
            file_type = path.suffix.lower()

            # Process based on file type
            content = self._extract_content(path)

            # Calculate word count
            word_count = len(content.split()) if content else 0

            # Create metadata
            metadata = {
                "file_path": str(path),
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "creation_time": file_stats.st_ctime,
                "modification_time": file_stats.st_mtime,
                "file_extension": file_type,
            }

            self.logger.info(f"Successfully processed {path.name}: {word_count} words extracted")

            return ProcessedDocument(
                filename=path.name,
                content=content,
                metadata=metadata,
                word_count=word_count,
                file_type=file_type,
                size_bytes=file_size,
                processing_status="success",
            )

        except Exception as e:
            self.logger.error(f"Error processing file {path.name}: {str(e)}")
            return ProcessedDocument(
                filename=path.name,
                content="",
                metadata={},
                word_count=0,
                file_type=path.suffix.lower(),
                size_bytes=path.stat().st_size if path.exists() else 0,
                processing_status="error",
                error_message=str(e),
            )

    def _extract_content(self, file_path: Path) -> str:
        """
        Extract content from file using docling.

        Args:
            file_path: Path to the file

        Returns:
            str: Extracted text content
        """
        try:
            # Import docling here to handle cases where it's not installed
            from docling.document_converter import DocumentConverter

            # Initialize converter
            converter = DocumentConverter()

            # Convert document
            result = converter.convert(str(file_path))

            # Extract text content
            if hasattr(result, "document") and hasattr(result.document, "export_to_markdown"):
                content = result.document.export_to_markdown()
            elif hasattr(result, "document") and hasattr(result.document, "text"):
                content = result.document.text
            else:
                # Fallback: try to get any text content
                content = str(result)

            return content

        except ImportError:
            self.logger.warning("Docling not available, falling back to basic text extraction")
            return self._fallback_extract_content(file_path)
        except Exception as e:
            self.logger.warning(f"Docling processing failed for {file_path.name}: {e}")
            return self._fallback_extract_content(file_path)

    def _fallback_extract_content(self, file_path: Path) -> str:
        """
        Fallback content extraction for when docling is not available.

        Args:
            file_path: Path to the file

        Returns:
            str: Extracted text content
        """
        file_type = file_path.suffix.lower()

        try:
            if file_type in [".txt", ".md"]:
                # Plain text files
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()

            elif file_type == ".pdf":
                # Try basic PDF extraction if available
                try:
                    import PyPDF2

                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page in reader.pages:
                            content += page.extract_text() + "\n"
                        return content
                except ImportError:
                    return f"PDF file detected but PyPDF2 not available for extraction: {file_path.name}"

            elif file_type in [".docx"]:
                # Try basic DOCX extraction if available
                try:
                    import docx

                    doc = docx.Document(file_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    return content
                except ImportError:
                    return f"DOCX file detected but python-docx not available for extraction: {file_path.name}"

            else:
                return f"File type {file_type} requires docling for proper extraction: {file_path.name}"

        except Exception as e:
            return f"Error in fallback extraction: {str(e)}"

    def process_multiple_files(self, file_paths: List[Union[str, Path]]) -> List[ProcessedDocument]:
        """
        Process multiple files and return results.

        Args:
            file_paths: List of file paths to process

        Returns:
            List[ProcessedDocument]: List of processed documents
        """
        results = []

        for file_path in file_paths:
            result = self.process_file(file_path)
            results.append(result)

        return results


class DatasetBuilder:
    """
    Builds training datasets from processed documents.
    """

    def __init__(self, logger_name: str = "app.data.dataset_builder"):
        """Initialize the dataset builder."""
        self.logger = logging.getLogger(logger_name)

    def build_training_dataset(self, documents: List[ProcessedDocument]) -> Dict[str, Any]:
        """
        Build a training dataset from processed documents.

        Args:
            documents: List of processed documents

        Returns:
            Dict containing dataset information and statistics
        """
        # Filter successful documents
        successful_docs = [doc for doc in documents if doc.processing_status == "success"]

        if not successful_docs:
            return {"status": "error", "message": "No successfully processed documents", "dataset": None, "statistics": {}}

        # Combine all content
        combined_content = []
        total_words = 0
        file_types = {}

        for doc in successful_docs:
            combined_content.append({"filename": doc.filename, "content": doc.content, "metadata": doc.metadata, "word_count": doc.word_count})

            total_words += doc.word_count

            # Track file types
            file_type = doc.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1

        # Create dataset statistics
        statistics = {
            "total_documents": len(successful_docs),
            "total_words": total_words,
            "average_words_per_doc": round(total_words / len(successful_docs), 2) if successful_docs else 0,
            "file_types": file_types,
            "failed_documents": len(documents) - len(successful_docs),
        }

        self.logger.info(f"Built dataset with {len(successful_docs)} documents, {total_words} total words")

        return {"status": "success", "dataset": combined_content, "statistics": statistics}
